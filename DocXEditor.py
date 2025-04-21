import os
import shutil
import zipfile
from datetime import datetime
from itertools import count
from docx import Document as BayooDocument

from lxml import etree
from itertools import count
from copy import deepcopy


NS_W  = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NS_REL = "http://schemas.openxmlformats.org/package/2006/relationships"
NSMAP = {"w": NS_W}

# ------------- helper namespaces --------------
NS_CM = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
NS_C  = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NSMAP_COMMENTS = {"w": NS_C}


class DocXEditor:
    # ---------- life‑cycle helpers ----------
    def __init__(self, docx_path: str, output_path: str):
        self.docx_path   = docx_path
        self.output_path = output_path
        self.temp_dir    = "_tmp_docx"

        self._unzip()
        self._id_counter = count(self._highest_existing_change_id() + 1)
        self.doc_tree = self._load_xml("word/document.xml")

    def _unzip(self):
        """
        Unzip the document into a temporary directory.

        Returns:
            None
        """
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)
        with zipfile.ZipFile(self.docx_path) as zf:
            zf.extractall(self.temp_dir)

    @property
    def paragraphs(self):
        """
        A live list of all <w:p> elements in the document.
        Accessing this property always re‑queries the current XML tree,
        so it automatically reflects any edits you make.
        """
        return BayooDocument(self.docx_path).paragraphs

    def _load_xml(self, relative_path: str) -> etree._ElementTree:
        """
        Load an XML file from the temp directory.

        Args:
            relative_path: The path to the file within the temp directory.

        Returns:
            The loaded XML tree.
        """
        full = os.path.join(self.temp_dir, relative_path)
        return etree.parse(full)

    def _write_xml(self, tree: etree._ElementTree, relative_path: str):
        """
        Write an XML tree to a file in the temp directory.

        Args:
            tree: The XML tree to write.
            relative_path: The path to the file within the temp directory.

        Returns:
            None
        """
        full = os.path.join(self.temp_dir, relative_path)
        tree.write(full, xml_declaration=True, encoding="UTF-8", standalone="yes")

    # ---------- tracked‑change ID helpers ----------
    def _highest_existing_change_id(self) -> int:
        """
        Find the highest existing tracked change ID in the document.

        Returns:
            The highest existing tracked change ID, or 0 if none exist.
        """
        try:
            tree = self._load_xml("word/document.xml")
        except FileNotFoundError:
            return 0
        ids = tree.xpath('//w:ins/@w:id | //w:del/@w:id', namespaces=NSMAP)
        return max(map(int, ids), default=0)

    def _next_change_id(self) -> str:
        """
        Get the next available tracked change ID.

        Returns:
            The next available tracked change ID.
        """
        return str(next(self._id_counter))

    def get_full_text(self, para):
        return "".join(t.text for t in para.xpath(".//w:t", namespaces=NSMAP) if t.text)

    def normalize(self, text):
        return re.sub(r'\s+', ' ', text.strip())

    # ---------- tracked change ----------
    def modify_text_in_doc(self, old_text, new_text, full_paragraph_text, author="ChatGPT"):
        """
        Modify text in the document by replacing occurrences of `old_text` with `new_text`.

        Args:
            old_text: The text to be replaced.
            new_text: The text to replace with.
            author: The author of the tracked change.

        Returns:
            None
        """
        for para in self.doc_tree.xpath("//w:p", namespaces=NSMAP):
            for t in para.xpath(".//w:t", namespaces=NSMAP):
                if not t.text or old_text not in t.text:
                    continue

                before, _, after = t.text.partition(old_text)
                parent_run = t.getparent()
                p = parent_run.getparent()
                run_idx = p.index(parent_run)

                # clear original run’s text (or remove it if empty)
                if before:
                    t.text = before
                    run_idx += 1
                else:
                    p.remove(parent_run)

                change_id = self._next_change_id()

                # deletion block
                del_el  = etree.Element(f"{{{NS_W}}}del",  nsmap=NSMAP,
                                        attrib={f"{{{NS_W}}}id": change_id,
                                                f"{{{NS_W}}}author": author,
                                                f"{{{NS_W}}}date": datetime.utcnow().isoformat()})
                del_run = etree.SubElement(del_el, f"{{{NS_W}}}r")
                del_txt = etree.SubElement(del_run, f"{{{NS_W}}}delText")
                del_txt.text = old_text
                p.insert(run_idx, del_el)
                run_idx += 1

                # insertion block
                ins_el  = etree.Element(f"{{{NS_W}}}ins", nsmap=NSMAP,
                                        attrib={f"{{{NS_W}}}id": change_id,
                                                f"{{{NS_W}}}author": author,
                                                f"{{{NS_W}}}date": datetime.utcnow().isoformat()})
                ins_run = etree.SubElement(ins_el, f"{{{NS_W}}}r")
                ins_txt = etree.SubElement(ins_run, f"{{{NS_W}}}t")
                ins_txt.text = " " + new_text
                p.insert(run_idx, ins_el)
                run_idx += 1

                # trailing text
                if after:
                    run_after = etree.Element(f"{{{NS_W}}}r")
                    t_after   = etree.SubElement(run_after, f"{{{NS_W}}}t")
                    t_after.text = after
                    p.insert(run_idx, run_after)
                break  # one match per paragraph


    def add_comment_to_paragraph(
            self,
            target_text: str,
            comment_text: str,
            author: str = "Reviewer",
            initials: str = "RV",
            *,
            new_paragraph: bool = False
    ):
        """
        Find the *first* paragraph that contains ``target_text``,
        split the run so that ``target_text`` becomes a stand‑alone run,
        then add an **empty** run (or a brand‑new paragraph) immediately
        after it and attach a Word comment to that blank run/paragraph.

        When ``new_paragraph`` is True we insert an *entire* paragraph
        after the original one; otherwise we stay in the same paragraph.
        """
        if not target_text or not target_text.strip():
            print("⚠️  Skipping comment: <original_text> is empty.")
            return

        def _insert_paragraph_after(para):
            """
            python‑docx has no helper for “insert‑after”, so we do it
            at the XML level and then wrap it as a Paragraph object.
            """
            new_p = deepcopy(para._p)          # shallow, empty copy
            for n in list(new_p):              # strip content
                new_p.remove(n)
            para._p.addnext(new_p)
            return para._parent._body._paragraphs[-1]   # last para in body

        doc   = BayooDocument(self.docx_path)
        found = False

        for para in doc.paragraphs:
            if target_text not in para.text:
                continue

            # ---- locate the run that *contains* the text -----------------
            run_idx = None
            for i, run in enumerate(para.runs):
                if target_text in run.text:
                    run_idx = i
                    break

            if run_idx is None:
                continue     # (shouldn’t happen)

            run = para.runs[run_idx]
            before, _, after = run.text.partition(target_text)

            # ---- split current run into before | target | after ----------
            # keep original formatting
            run.text = before

            target_run = para.add_run(target_text)
            target_run.style = run.style

            after_run = None
            if after:
                after_run = para.add_run(after)
                after_run.style = run.style

            # re‑order so they appear *immediately* after the “before” run
            # (python‑docx appends by default)
            run._r.addnext(target_run._r)
            if after_run is not None:
                target_run._r.addnext(after_run._r)

            # ---- decide where the comment anchor will live ---------------
            if new_paragraph:
                anchor_para = _insert_paragraph_after(para)
                anchor_run  = anchor_para.add_run("")      # truly blank
            else:
                anchor_run  = para.add_run("")             # blank run
                target_run._r.addnext(anchor_run._r)       # keep order

            # ---- finally add the comment --------------------------------
            anchor_run.add_comment(
                text     = comment_text,
                author   = author,
                initials = initials
            )
            found = True
            break

        if not found:
            print("❌  Text not found – no comment added.")
            return

        # ---- persist, reload internal state ------------------------------
        doc.save(self.output_path)

        self.docx_path = self.output_path
        self._unzip()                                  # refresh temp dir
        self._id_counter = count(self._highest_existing_change_id() + 1)
        self.doc_tree    = self._load_xml("word/document.xml")
        print("✅  Comment added and file saved.")


    def save(self):
        """
        Write back document.xml and zip everything.
        
        Returns:
            None
        """
        self._write_xml(self.doc_tree, "word/document.xml")

        # zip everything
        with zipfile.ZipFile(self.output_path, "w", zipfile.ZIP_DEFLATED) as zf:
            for root, _, files in os.walk(self.temp_dir):
                for f in files:
                    full = os.path.join(root, f)
                    arc  = os.path.relpath(full, self.temp_dir)
                    zf.write(full, arc)

        shutil.rmtree(self.temp_dir, ignore_errors=True)

if __name__ == "__main__":
    doc = DocXEditor("input.docx", "output_fixed.docx")
    # Apply all comments first
    doc.add_comment_to_paragraph("INTRODUCTION", "Please clarify what this refers to.")
    doc.modify_text_in_doc("INTRODUCTION", "FUAD WAS HERE", author="Fuad")
    doc.save()
    